"""
Modality 1: DOCUMENT
Extracts raw text from uploaded resume files (PDF / DOCX / TXT).
This is real binary-file parsing, not just reading strings.
"""
import os
import pdfplumber
import docx


class ParseError(Exception):
    pass


def extract_text(file_path: str) -> str:
    """Detect file type by extension and extract plain text from it."""
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return _extract_pdf(file_path)
    if ext == ".docx":
        return _extract_docx(file_path)
    if ext == ".txt":
        return _extract_txt(file_path)

    raise ParseError(f"Unsupported file type: {ext}. Use PDF, DOCX, or TXT.")


def _extract_pdf(file_path: str) -> str:
    text_chunks = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            text_chunks.append(page_text)
    text = "\n".join(text_chunks).strip()
    if not text:
        raise ParseError(
            "Could not extract any text from this PDF (it may be a scanned "
            "image without a text layer)."
        )
    return text


def _extract_docx(file_path: str) -> str:
    document = docx.Document(file_path)
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    text = "\n".join(parts).strip()
    if not text:
        raise ParseError("Could not extract any text from this DOCX file.")
    return text


def _extract_txt(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read().strip()
    if not text:
        raise ParseError("This TXT file appears to be empty.")
    return text
